import os
import pandas as pd
import pdfplumber
import docx
import re
from datetime import datetime


def extract_text_and_title_from_pdf(file_path):
    """
    Extracts text and identifies a potential title from a PDF file.
    """
    text = ""
    title = None

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
            
            # Extract title: Find the largest text element on the page
            largest_font = 0
            for char in page.chars:
                font_size = char.get("size", 0)
                if font_size > largest_font:
                    largest_font = font_size
                    title = char.get("text", "").strip()
    
    return text.strip(), title


def extract_text_and_title_from_docx(file_path):
    """
    Extracts text and identifies a potential title from a DOCX file.
    """
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    title = None
    
    # Assume title is in the first non-empty paragraph
    for para in doc.paragraphs:
        if para.text.strip():
            title = para.text.strip()
            break

    return text, title


def extract_fields(content, title):
    """
    Extracts name, student ID, and validates the title based on conditions.
    """
    # Match student ID based on markers like "ID:" or "BinusianID:"
    student_id_pattern = r"(?:ID:|BinusianID:)\s*(\d+)"
    student_id_match = re.search(student_id_pattern, content)
    student_id = student_id_match.group(1).strip() if student_id_match else None

    # Extract name if possible
    name_pattern = r"([A-Za-z\s]+)\s*"
    name_match = re.search(name_pattern, content)
    name = name_match.group(1).strip() if name_match else None

    # Validate title if passed
    if not title:
        paragraphs = content.split("\n")
        title = next((para.strip() for para in paragraphs if para.strip()), None)

    return {
        "name": name,
        "student_id": student_id,
        "title": title
    }


def process_folder(directory):
    """
    Processes all files in the given directory and extracts relevant fields.
    """
    data = []
    for root, _, files in os.walk(directory):
        for file in files:
            path = os.path.join(root, file)
            extension = os.path.splitext(file)[-1].lower()
            content = ""
            title = None

            try:
                if extension == ".pdf":
                    content, title = extract_text_and_title_from_pdf(path)
                elif extension == ".docx":
                    content, title = extract_text_and_title_from_docx(path)
                else:
                    continue 
                
                fields = extract_fields(content, title)
                timestamp = os.path.getctime(path)
                readable_date = datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M:%S')
                
                data.append({
                    "filename": file,
                    "extension": extension,
                    "size": os.path.getsize(path),
                    "created_at": readable_date,
                    "content": content,
                    "name": fields["name"],
                    "title": fields["title"],
                    "student_id": fields["student_id"]
                })
            except Exception as e:
                print(f"Error processing {file}: {e}")
    
    return pd.DataFrame(data)


if __name__ == "__main__":
    directory = input("Enter the directory path containing files: ")
    processed_data = process_folder(directory)
    processed_data.to_csv('processed_data_with_fields.csv', index=False)
    print("Data preparation complete. Saved to 'processed_data_with_fields.csv'.")
